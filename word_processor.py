"""Word document processing"""
import docx
from PIL import Image
from io import BytesIO
import base64
from pathlib import Path
from langchain_core.documents import Document  # ← FIXED THIS LINE
from typing import List, Dict, Tuple
import logging

logger = logging.getLogger(__name__)


class WordProcessor:
    """Process Word documents"""
    
    def __init__(self):
        self.image_store = {}
    
    def process_word(self, docx_path: Path) -> List[Document]:
        """
        Extract text and images from Word document
        
        Returns:
            List of Document objects
        """
        doc = docx.Document(docx_path)
        documents = []
        
        # Extract text
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        if full_text:
            text_content = "\n\n".join(full_text)
            text_doc = Document(
                page_content=text_content,
                metadata={'page': 0, 'type': 'text', 'source': str(docx_path)}
            )
            documents.append(text_doc)
        
        # Extract images
        img_index = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_blob = rel.target_part.blob
                    pil_image = Image.open(BytesIO(image_blob)).convert("RGB")
                    
                    # Create unique ID
                    image_id = f"word_img_{img_index}"
                    
                    # Store as base64
                    buffered = BytesIO()
                    pil_image.save(buffered, format="PNG")
                    img_base64 = base64.b64encode(buffered.getvalue()).decode()
                    self.image_store[image_id] = img_base64
                    
                    # Create document
                    img_doc = Document(
                        page_content=f"[Image: {image_id}]",
                        metadata={
                            'page': 0,
                            'type': 'image',
                            'image_id': image_id,
                            'image_obj': pil_image,
                            'source': str(docx_path)
                        }
                    )
                    documents.append(img_doc)
                    img_index += 1
                    
                except Exception as e:
                    logger.warning(f"Could not extract image: {e}")
        
        return documents
